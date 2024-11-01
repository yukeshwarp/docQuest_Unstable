import streamlit as st
import json
from utils.pdf_processing import process_pdf_pages
from utils.llm_interaction import ask_question
from concurrent.futures import ThreadPoolExecutor, as_completed
import logging
import io
from docx import Document
import tiktoken

def count_tokens(text, model="gpt-4o"):
    encoding = tiktoken.encoding_for_model(model)
    tokens = encoding.encode(text)
    return len(tokens)

doc_token = 0
# Initialize session state
if 'documents' not in st.session_state:
    st.session_state.documents = {}
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'uploaded_files' not in st.session_state:
    st.session_state.uploaded_files = []

# Handle user question
def handle_question(prompt):
    if prompt:
        try:
            with st.spinner('Thinking...'):
                answer, tot_tokens = ask_question(
                    st.session_state.documents, prompt, st.session_state.chat_history
                )
            
            st.session_state.chat_history.append({
                "question": prompt,
                "answer": f"{answer}\nTotal tokens: {tot_tokens}"
            })

        except Exception as e:
            st.error(f"Error processing question: {e}")

# Reset session state
def reset_session():
    st.session_state.documents = {}
    st.session_state.chat_history = []
    st.session_state.uploaded_files = []

# Function to display chat history
def display_chat():
    if st.session_state.chat_history:
        for i, chat in enumerate(st.session_state.chat_history):
            user_message = f"""
            <div style='padding:10px; border-radius:10px; margin:5px 0; text-align:right;'> 
            {chat['question']}
            </div>
            """
            assistant_message = f"""
            <div style='padding:10px; border-radius:10px; margin:5px 0; text-align:left;'> 
            {chat['answer']}
            </div>
            """
            st.markdown(user_message, unsafe_allow_html=True)
            st.markdown(assistant_message, unsafe_allow_html=True)

            # Prepare content for download
            chat_content = {
                "question": chat["question"],
                "answer": chat["answer"]
            }

            # Generate Word document for download
            def generate_word_document(content):
                doc = Document()
                doc.add_heading('Chat Response', 0)
                doc.add_paragraph(f"Question: {content['question']}")
                doc.add_paragraph(f"Answer: {content['answer']}")
                return doc

            # Generate the Word document in memory
            doc = generate_word_document(chat_content)
            word_io = io.BytesIO()
            doc.save(word_io)
            word_io.seek(0)

            # Provide a download button for the Word document
            st.download_button(
                label="📥",
                data=word_io,
                file_name=f"chat_{i+1}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Sidebar for file upload
with st.sidebar:
    uploaded_files = st.file_uploader(
        " ",
        type=["pdf", "docx", "xlsx", "pptx"],
        accept_multiple_files=True,
        help="Supports PDF, DOCX, XLSX, and PPTX formats.",
    )

    if uploaded_files:
        new_files = []
        for index, uploaded_file in enumerate(uploaded_files):
            if uploaded_file.name not in st.session_state.documents:
                new_files.append(uploaded_file)
            else:
                st.info(f"{uploaded_file.name} is already uploaded.")

        if new_files:
            progress_text = st.empty()
            progress_bar = st.progress(0)
            total_files = len(new_files)

            with st.spinner("Learning about your document(s)..."):
                with ThreadPoolExecutor(max_workers=2) as executor:
                    future_to_file = {
                        executor.submit(process_pdf_pages, uploaded_file, first_file=(index == 0)): uploaded_file 
                        for index, uploaded_file in enumerate(new_files)
                    }

                    for i, future in enumerate(as_completed(future_to_file)):
                        uploaded_file = future_to_file[future]
                        try:
                            document_data = future.result()
                            doc_token += count_tokens(str(document_data))
                            st.session_state.documents[uploaded_file.name] = document_data
                            st.success(f"{uploaded_file.name} processed successfully!")
                        except Exception as e:
                            st.error(f"Error processing {uploaded_file.name}: {e}")

                        progress_bar.progress((i + 1) / total_files)

            st.sidebar.write(f"Total document tokens: {doc_token}")
            progress_text.text("Processing complete.")
            progress_bar.empty()

    if st.session_state.documents:
        download_data = json.dumps(st.session_state.documents, indent=4)
        st.download_button(
            label="Download Document Analysis",
            data=download_data,
            file_name="document_analysis.json",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# App header and main interface
st.image("logoD.png", width=200)
st.title("docQuest")
st.subheader("Unveil the Essence, Compare Easily, Analyze Smartly", divider="orange")

# Chat interface
if st.session_state.documents:
    prompt = st.chat_input("Ask me anything about your documents", key="chat_input")
    if prompt:
        handle_question(prompt)  

# Display the chat history
display_chat()
